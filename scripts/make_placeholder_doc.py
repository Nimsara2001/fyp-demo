"""Generate one tiny, self-authored Sinhala .docx for pipeline wiring smoke-tests.

This is NOT sourced from akshara-kit's own tests/fixtures/* (a separate, unresolved
copyright concern) and is not meant to demonstrate anything about extraction quality —
it exists purely so both ingestion pipelines have something to run against before real
sample documents arrive. Delete it once real documents are in data/raw/.
"""

from __future__ import annotations

from docx import Document

from fyp_demo import config

PARAGRAPHS = [
    "මෙය අකුරු කට්ටලය පරීක්ෂා කිරීම සඳහා සකස් කරන ලද නියැදි ලේඛනයකි.",
    "මෙම ලේඛනයේ අරමුණ වන්නේ සිංහල පෙළ නිසි ලෙස කියවීම හා වර්ගීකරණය කළ හැකි දැයි පරීක්ෂා කිරීමයි.",
    "පර්යේෂණ ව්‍යාපෘතියක් සඳහා විවිධ ලේඛන වර්ග විශ්ලේෂණය කිරීම අවශ්‍ය වේ.",
    "පළමු පියවර ලෙස ලේඛනයෙන් පෙළ උපුටා ගැනීම සිදු කරනු ලැබේ.",
    "දෙවන පියවර ලෙස උපුටාගත් පෙළ කුඩා කොටස් වලට වෙන් කරනු ලැබේ.",
    "අවසාන පියවර ලෙස එම කොටස් සෙවුම් පද්ධතියක් තුළට ඇතුළත් කරනු ලැබේ.",
]


def build_placeholder_docx(path=None):
    path = path or (config.DATA_RAW_DIR / "placeholder.docx")
    path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading("නියැදි ලේඛනය", level=1)
    for paragraph in PARAGRAPHS:
        doc.add_paragraph(paragraph)

    doc.save(path)
    return path


if __name__ == "__main__":
    saved_path = build_placeholder_docx()
    print(f"Wrote placeholder document to {saved_path}")
